"""
CRUD Operations with CosmosDB Integration
Handles document upload, indexing, and querying using CosmosDB storage
"""
import os
import re
import tempfile
from pathlib import Path
from fastapi import UploadFile
from app.model import Document
from app.cosmos_input import InputManager
from app.cosmos_output import OutputManager
from app.cosmos_logs import LogsManager
from app.cosmos_config import config_manager
# from app.cosmos_db import get_next_document_id  # No longer needed - using firm-specific managers
from graphrag.config.create_graphrag_config import create_graphrag_config
from graphrag.config.enums import IndexingMethod
from PyPDF2 import PdfReader
from pdf2image import convert_from_path
import pytesseract
import yaml
import pandas as pd
from graphrag.api.query import global_search, local_search, drift_search
from graphrag.api.index import build_index
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.cosmos_client import CosmosDBManager
RAG_ROOT = "ragtest"
CONFIG_PATH = os.path.join(RAG_ROOT, "settings.yaml").replace("\\", "/")
POPPLER_PATH = os.getenv("POPPLER_PATH", r"C:\poppler-25.07.0\Library\bin")
TESSERACT_PATH = os.getenv("TESSERACT_PATH", r"C:\Program Files\Tesseract-OCR")

pytesseract.pytesseract.tesseract_cmd = os.path.join(TESSERACT_PATH, "tesseract.exe")

# Initialize managers (default - will be overridden per firm)
input_manager = InputManager()
output_manager = OutputManager()
logs_manager = LogsManager()

# Cache for firm managers to avoid re-initialization
_firm_managers_cache = {}

def get_firm_managers(firm_id: str, case_id: str = None):
    """Get or create firm-specific managers with case-specific containers"""
    if case_id is None:
        raise ValueError("case_id is required for case-specific containers")
    
    cache_key = f"{firm_id}_{case_id}"
    
    if cache_key not in _firm_managers_cache:
        # Create firm-specific database manager
        firm_manager = CosmosDBManager()
        firm_manager.database_name = f"graphrag_{firm_id}"
        firm_manager.initialize_database()
        
        # Create case-specific containers
        firm_manager.create_case_specific_containers(case_id)
        
        # Create firm-specific managers
        firm_input_manager = InputManager()
        firm_input_manager.manager = firm_manager
        firm_input_manager.container = firm_manager.input_container
        
        firm_output_manager = OutputManager()
        firm_output_manager.manager = firm_manager
        firm_output_manager.container = firm_manager.output_container
        
        # Initialize LogsManager
        firm_logs_manager = LogsManager()
        firm_logs_manager.manager = firm_manager
        firm_logs_manager.container = firm_manager.logs_container
        
        _firm_managers_cache[cache_key] = {
            'input': firm_input_manager,
            'output': firm_output_manager,
            'logs': firm_logs_manager
        }
    
    return _firm_managers_cache[cache_key]
# ==================== Text Extraction ====================

def clean_ocr_text(text: str) -> str:
    """Clean OCR text."""
    text = re.sub(r"_+", " ", text)
    text = re.sub(r"\s{2,}", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX files."""
    try:
        doc = DocxDocument(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        return '\n'.join(text)
    except Exception as e:
        print(f"Error extracting text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_xlsx(file_path: str) -> str:
    """Extract text from XLSX files."""
    try:
        workbook = load_workbook(file_path, data_only=True)
        text = []
        
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            text.append(f"Sheet: {sheet_name}")
            
            for row in sheet.iter_rows(values_only=True):
                row_text = []
                for cell in row:
                    if cell is not None:
                        row_text.append(str(cell))
                if row_text:
                    text.append('\t'.join(row_text))
            text.append('')
        
        return '\n'.join(text)
    except Exception as e:
        print(f"Error extracting text from XLSX {file_path}: {e}")
        return ""


# ==================== Document Upload ====================

async def save_document(file: UploadFile, case_id: str, firm_id: str) -> Document:
    """
    Save a document to CosmosDB with text extraction.
    
    Args:
        file: The uploaded file
        case_id: The case identifier
        firm_id: The firm identifier
        
    Returns:
        Document object with metadata
    """
    print(f"[INFO] ===== save_document called =====")
    print(f"[INFO] File: {file.filename}, Firm ID: {firm_id}, Case ID: {case_id}")
    
    # Get firm-specific managers
    firm_managers = get_firm_managers(firm_id, case_id)
    firm_input_manager = firm_managers['input']
    firm_output_manager = firm_managers['output']
    firm_logs_manager = firm_managers['logs']    
    try:
        firm_logs_manager.info(case_id, f"Starting document upload: {file.filename}", "upload")
    except Exception as log_err:
        print(f"[WARNING] Could not log to CosmosDB: {log_err}")
    
    new_doc_id = firm_input_manager.get_next_document_id(case_id)
    file_extension = os.path.splitext(file.filename)[1].lower()
    new_filename = f"{firm_id}_{case_id}_{new_doc_id}{file_extension}"
    
    print(f"[INFO] Saving document: {file.filename} -> {new_filename}")
    
    try:
        firm_logs_manager.info(case_id, f"Saving document: {file.filename} -> {new_filename}", "upload")
    except Exception as log_err:
        print(f"[WARNING] Could not log to CosmosDB: {log_err}")
    
    # Read file content
    file_content = await file.read()
    
    # NOTE: We do NOT store the original PDF/DOCX binary in CosmosDB
    # Only the extracted text is stored (saves storage space + GraphRAG only needs text)
    print(f"   [INFO] File loaded into memory: {len(file_content)} bytes")
    
    # Extract text based on file type
    extracted_text = ""
    
    # Save to temporary file for processing
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as tmp_file:
        tmp_file.write(file_content)
        tmp_file_path = tmp_file.name
    
    try:
        if file_extension == ".txt":
            # Handle plain text files
            print(f"[INFO] Reading text file: {new_filename}")
            firm_logs_manager.info(case_id, f"Reading text file: {new_filename}", "upload")
            extracted_text = file_content.decode('utf-8', errors='ignore')
            print(f"   [SUCCESS] Text file read: {len(extracted_text)} characters")
        
        elif file_extension == ".pdf":
            print(f"[INFO] Extracting text from PDF: {new_filename}")
            firm_logs_manager.info(case_id, f"Extracting text from PDF: {new_filename}", "upload")
            
            try:
                reader = PdfReader(tmp_file_path)
                extracted_text = "".join(page.extract_text() or "" for page in reader.pages)
                print(f"   [SUCCESS] PDF text extraction successful: {len(extracted_text)} characters")
            except Exception as e:
                print(f"   [ERROR] PDF text extraction failed: {e}")
                logs_manager.error(case_id, f"PDF text extraction failed: {e}", "upload")
            
            # Fallback to OCR if no text
            if not extracted_text.strip():
                print(f"   [INFO] No text found, running OCR on {new_filename}...")
                firm_logs_manager.info(case_id, f"Running OCR on {new_filename}", "upload")
                
                try:
                    # Primary OCR attempt using pdf2image + pytesseract
                    images = convert_from_path(tmp_file_path, poppler_path=r"C:\poppler-25.07.0\Library\bin")
                    ocr_text = []
                    for img in images:
                        text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
                        ocr_text.append(text)
                    extracted_text = "\n".join(ocr_text)
                    if extracted_text.strip():
                        print(f"   [SUCCESS] OCR successful: {len(extracted_text)} characters")
                    else:
                        print(f"   [WARNING] OCR returned empty text for {new_filename}")
                        firm_logs_manager.warning(case_id, f"OCR produced empty text for {new_filename}", "upload")

                except Exception as ocr_e:
                    print(f"   [ERROR] OCR failed for {new_filename}: {ocr_e}")
                    firm_logs_manager.error(case_id, f"OCR failed: {ocr_e}", "upload")

                    # ✅ Fallback using PyMuPDF (if installed)
                    try:
                        import fitz  # PyMuPDF
                        from PIL import Image
                        print("   [INFO] Retrying OCR using PyMuPDF...")
                        with fitz.open(tmp_file_path) as doc:
                            ocr_text = []
                            for page in doc:
                                pix = page.get_pixmap()
                                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                                text = pytesseract.image_to_string(img, lang="eng", config="--psm 6")
                                ocr_text.append(text)
                        extracted_text = "\n".join(ocr_text)
                        print(f"   [SUCCESS] Fallback OCR extracted {len(extracted_text)} characters")
                    except Exception as fallback_e:
                        print(f"   [ERROR] Fallback OCR also failed: {fallback_e}")
                        firm_logs_manager.error(case_id, f"Fallback OCR failed: {fallback_e}", "upload")
                        extracted_text = ""
        
        elif file_extension == ".docx":
            print(f"[INFO] Extracting text from DOCX: {new_filename}")
            firm_logs_manager.info(case_id, f"Extracting text from DOCX: {new_filename}", "upload")
            extracted_text = extract_text_from_docx(tmp_file_path)
            print(f"   [SUCCESS] DOCX text extraction: {len(extracted_text)} characters")
        
        elif file_extension == ".xlsx":
            print(f"[INFO] Extracting text from XLSX: {new_filename}")
            firm_logs_manager.info(case_id, f"Extracting text from XLSX: {new_filename}", "upload")
            extracted_text = extract_text_from_xlsx(tmp_file_path)
            print(f"   [SUCCESS] XLSX text extraction: {len(extracted_text)} characters")
    
    finally:
        # Clean up temporary file
        if os.path.exists(tmp_file_path):
            os.unlink(tmp_file_path)
    
    # Clean and store extracted text
    # Default to original filename if no text extracted
    stored_filename = new_filename
    
    if extracted_text:
        extracted_text = clean_ocr_text(extracted_text)
        # Use .txt extension for plain text content
        txt_filename = f"{firm_id}_{case_id}_{new_doc_id}.txt"
        stored_filename = txt_filename  # Update to use txt filename
        
        # getting uploaded the original text to cosmosDB
        firm_input_manager.store_extracted_text(
            case_id=case_id,
            document_id=new_doc_id,
            text_content=extracted_text,
            filename=txt_filename,
            original_filename=file.filename,  # Track original filename
            firm_id=firm_id  # Pass firm_id for proper database routing
        )
        
        print(f"   [SUCCESS] Text extracted and saved to CosmosDB: {txt_filename} ({len(extracted_text)} characters)")
        print(f"   [INFO] Original PDF/DOCX binary NOT stored (only text is kept)")
        try:
            firm_logs_manager.info(case_id, f"Text saved: {txt_filename} from {file.filename} ({len(extracted_text)} chars)", "upload")
        except Exception as log_err:
            print(f"[WARNING] Could not log to CosmosDB: {log_err}")
    else:
        print(f"   [WARNING] No text could be extracted from {new_filename}")
        firm_logs_manager.warning(case_id, f"No text extracted from {new_filename}", "upload")
    
    return Document(
        firm_id=firm_id,
        case_id=case_id,
        document_id=new_doc_id,
        filename=stored_filename,  # Return the actual stored filename (.json)
        file_path=f"cosmos://{case_id}/{new_doc_id}"
    )


# ==================== Indexing ====================

async def index_documents_for_case(case_id: str, firm_id: str):
    """
    Index documents for a case using GraphRAG with native CosmosDB storage.
    
    This function:
    1. Checks if containers exist - if yes, sets is_update_run=True for incremental indexing
    2. If containers don't exist, creates them
    3. Loads configuration from CosmosDB
    4. Configures GraphRAG to read directly from CosmosDB input container
    5. Runs GraphRAG indexing (native CosmosDB read/write)
    6. Stores all results (graph, vectors, cache) directly to CosmosDB
    7. For incremental indexing: merges entities/relationships from update_output to output
    8. Marks documents as indexed
    """
    # Create firm-specific database manager to check containers
    firm_manager = CosmosDBManager()
    firm_manager.database_name = f"graphrag_{firm_id}"
    firm_manager.initialize_database()
    
    # Check if containers exist AND have indexed data - if yes, this is an incremental update
    # This check verifies that output container has final entities (UUID IDs), not just raw data
    print(f"[CHECK] Checking if case {case_id} has previous indexed data...")
    has_indexed_data = firm_manager.check_case_containers_exist(case_id)
    
    # Create containers if they don't exist
    if not has_indexed_data:
        print(f"[INFO] ✓ First run detected for case {case_id}")
        print(f"[INFO] Creating containers and running FULL indexing (is_update_run=False)")
        firm_manager.create_case_specific_containers(case_id)
        is_update_run = False  # Ensure it's set to False for first run
    else:
        # For incremental indexing, delta and previous containers will be created lazily
        print(f"[INFO] ✓ Previous indexing found for case {case_id}")
        print(f"[INFO] Running INCREMENTAL indexing (is_update_run=True)")
        print(f"[INFO] Delta and previous containers will be created automatically when data is written")
        # Get existing containers (this won't recreate them, just gets references)
        firm_manager.create_case_specific_containers(case_id)
        is_update_run = True  # Ensure it's set to True for incremental run
    
    print(f"[CONFIRM] Indexing mode for case {case_id}: {'INCREMENTAL' if is_update_run else 'FULL'}")
    
    # Get firm-specific managers
    firm_managers = get_firm_managers(firm_id, case_id)
    firm_input_manager = firm_managers['input']
    firm_output_manager = firm_managers['output']
    firm_logs_manager = firm_managers['logs']      
    
    log_message = "Starting incremental indexing process" if is_update_run else "Starting indexing process"
    firm_logs_manager.info(case_id, log_message, "indexing")
    
    # GraphRAG logs will be saved to ragtest/logs/{case_id}/ (persistent)
    persistent_logs_dir = Path(RAG_ROOT).resolve() / "logs" / case_id
    persistent_logs_dir.mkdir(parents=True, exist_ok=True)
    
    # Load prompts from local files (no CosmosDB needed)
    print(f"[SETUP] Loading prompts from local files...")
    prompts = config_manager.get_all_prompts()
    if prompts:
        print(f"[SUCCESS] Loaded {len(prompts)} prompts from local files")
    else:
        print(f"[WARNING] No prompts found in local directory")
    
    try:
        # GraphRAG will read directly from CosmosDB (using your custom _read_with_flexible_partition)
        print(f"[COSMOS] GraphRAG will read directly from CosmosDB input container for case {case_id}")
        firm_logs_manager.info(case_id, "GraphRAG reading from CosmosDB", "indexing")
        
        # Load settings.yaml from local file (no CosmosDB needed)
        print(f"[CONFIG] Loading settings.yaml from local file...")
        config_content = config_manager.get_settings_yaml()
        
        if not config_content:
            print(f"[ERROR] settings.yaml not found in local file")
            raise FileNotFoundError("settings.yaml not found in ragtest directory")
        else:
            print(f"[SUCCESS] Loaded settings.yaml from local file")
        
        # Expand environment variables
        for var in os.environ:
            config_content = config_content.replace(f"${{{var}}}", os.environ[var])
        
        config_dict = yaml.safe_load(config_content)
        
        # Ensure all required config sections exist
        if "input" not in config_dict:
            raise ValueError("Config missing 'input' section")
        if "storage" not in config_dict["input"]:
            config_dict["input"]["storage"] = {}
        if "file_filter" not in config_dict["input"]:
            config_dict["input"]["file_filter"] = {}
        
        # Set case-specific filters for CosmosDB (GraphRAG reads directly)
        config_dict["input"]["file_filter"]["case_id"] = case_id
        config_dict["input"]["file_filter"]["is_text"] = "true"  # Must match string "true" in documents
        
        # For CosmosDB storage, base_dir is used as the database name
        # Get database name from environment or use default
        database_name = f"graphrag_{firm_id}"

        # ===== FIX: Configure database names for different storage types =====
        # The container names are already specified in settings.yaml
        # We only need to set the correct database names (base_dir)
        
        # Ensure output, cache, reporting, and update_index_output sections exist
        if "output" not in config_dict:
            config_dict["output"] = {}
        if "cache" not in config_dict:
            config_dict["cache"] = {}
        if "reporting" not in config_dict:
            config_dict["reporting"] = {}
        if "update_index_output" not in config_dict:
            config_dict["update_index_output"] = {}

        # Set database names for all storage types
        config_dict["input"]["storage"]["base_dir"] = database_name
        config_dict["output"]["base_dir"] = database_name
        config_dict["cache"]["base_dir"] = database_name
        
        # Get connection string for vector store and update_index_output
        connection_string = os.getenv("COSMOS_CONNECTION_STRING")
        if not connection_string:
            endpoint = os.getenv("COSMOS_ENDPOINT")
            key = os.getenv("COSMOS_KEY")
            if endpoint and key:
                connection_string = f"AccountEndpoint={endpoint};AccountKey={key};"
        if not connection_string:
            raise ValueError("COSMOS_CONNECTION_STRING or COSMOS_ENDPOINT+COSMOS_KEY must be set for indexing")
        
        # Configure update_index_output for CosmosDB (used ONLY during incremental indexing)
        if is_update_run:
            config_dict["update_index_output"]["type"] = "cosmosdb"
            config_dict["update_index_output"]["connection_string"] = connection_string
            config_dict["update_index_output"]["container_name"] = f"update_output_{case_id}"
            config_dict["update_index_output"]["base_dir"] = database_name
            config_dict["update_index_output"]["cosmosdb_account_blob_url"] = os.getenv("COSMOS_ENDPOINT")
            print(f"[CONFIG] Incremental indexing mode: Update index output container: {config_dict['update_index_output']['container_name']} (database: {database_name})")
        else:
            # For full indexing, update_index_output should not be used
            # Keep it as file type or leave it unconfigured
            print(f"[CONFIG] Full indexing mode: Update index output not needed")

        # Ensure vector store uses correct database name
        if "vector_store" not in config_dict:
            config_dict["vector_store"] = {}
        if "default_vector_store" not in config_dict["vector_store"]:
            config_dict["vector_store"]["default_vector_store"] = {}

        config_dict["vector_store"]["default_vector_store"]["database_name"] = database_name
        
        # Explicitly set connection_string for vector store to avoid AAD fallback during embedding
        # (connection_string was already set above, but ensure it's set for vector store)
        config_dict["vector_store"]["default_vector_store"]["connection_string"] = connection_string
        
        # Use persistent logs directory (ragtest/logs/{case_id})
        config_dict["reporting"]["base_dir"] = str(persistent_logs_dir)
        
        # ===== UPDATE: Set case-specific container names =====
        config_dict["input"]["storage"]["container_name"] = f"input_{case_id}"
        config_dict["output"]["container_name"] = f"output_{case_id}"
        config_dict["cache"]["container_name"] = f"cache_{case_id}"
        # Use case-specific base container directly; avoid runtime patching
        config_dict["vector_store"]["default_vector_store"]["container_name"] = f"output_{case_id}"
        
        # Add case_id to vector store config for our custom vector store
        config_dict["vector_store"]["default_vector_store"]["case_id"] = case_id
        
        print(f"[CONFIG] CosmosDB native configuration:")
        print(f"   Input:  {config_dict['input']['storage']['container_name']} (filter: case_id={case_id})")
        print(f"   Output: {config_dict['output']['container_name']} (database: {database_name})")
        print(f"   Cache:  {config_dict['cache']['container_name']} (database: {database_name})")
        print(f"   Vectors: {config_dict['vector_store']['default_vector_store']['container_name']} (database: {database_name})")
        print(f"   Logs:   {config_dict['reporting']['base_dir']} (persistent)")
        print(f"[INFO] GraphRAG will read/write directly from/to CosmosDB")
        
        # Note: No runtime patching of create_index_name during indexing; container base is case-specific
        
        # Create GraphRAG config with proper root directory for resolving relative paths
        # RAG_ROOT allows GraphRAG to resolve prompt paths like "prompts/extract_graph.txt"
        config_obj = create_graphrag_config(config_dict, RAG_ROOT)
        
        # Run GraphRAG indexing (reads/writes directly from/to CosmosDB)
        indexing_type = "incremental indexing" if is_update_run else "full indexing"
        print(f"[START] Starting GraphRAG {indexing_type} with native CosmosDB...")
        firm_logs_manager.info(case_id, f"Running GraphRAG {indexing_type} (native CosmosDB)", "indexing")
        
        await build_index(config=config_obj, method=IndexingMethod.Standard, is_update_run=is_update_run)
        
        print(f"[DONE] GraphRAG {indexing_type} completed for case {case_id}.")
        print(f"[INFO] All data stored directly in CosmosDB by GraphRAG")
        firm_logs_manager.info(case_id, f"GraphRAG {indexing_type} completed", "indexing")
        
        # For incremental indexing: merge entities and relationships from update_output to output
        if is_update_run:
            print(f"[INFO] Cleaning up raw data artifacts from incremental indexing...")
            firm_logs_manager.info(case_id, "Cleaning up raw data after incremental indexing", "indexing")
            
            # GraphRAG update workflows automatically copy final data to output container
            # We only need to clean up raw (numeric ID) entities, relationships, and embeddings
            cleanup_result = firm_output_manager.cleanup_raw_data_after_incremental_update(case_id, firm_id)
            
            print(f"[SUCCESS] Cleanup completed: {cleanup_result}")
            firm_logs_manager.info(case_id, f"Cleanup completed: {cleanup_result.get('status', 'unknown')}", "indexing")
        else:
            # Mark all documents as indexed (for full indexing)
            text_docs = firm_input_manager.list_documents(case_id, text_only=True)
            for doc in text_docs:
                firm_input_manager.mark_as_indexed(case_id, doc["document_id"])
            
            # Remove duplicate raw entities and relationships
            dedup_result = firm_output_manager.deduplicate_raw_data(case_id)
            print(f"[INFO] Deduplication result: {dedup_result}")
        
        print(f"[SUCCESS] All data uploaded to CosmosDB successfully!")
        firm_logs_manager.info(case_id, "Indexing completed successfully", "indexing")
    
    except Exception as e:
        print(f"[ERROR] GraphRAG indexing failed: {e}")
        firm_logs_manager.error(case_id, f"Indexing failed: {e}", "indexing")
        raise RuntimeError(f"Indexing failed: {e}")
    
    print(f"[SUCCESS] Indexing process completed successfully for case {case_id}.")


# ==================== Query ====================

async def get_answer(query_text: str, case_id: str, firm_id: str):
    """
    Query documents for a given case_id using GraphRAG with CosmosDB storage.
    """
    # Get firm-specific managers
    firm_managers = get_firm_managers(firm_id, case_id)
    firm_input_manager = firm_managers['input']
    firm_output_manager = firm_managers['output']
    firm_logs_manager = firm_managers['logs']      
    try:
        firm_logs_manager.info(case_id, f"Query: {query_text}", "query")
        
        # Load config
        with open(CONFIG_PATH, "r", encoding="utf-8") as f:
            config_content = f.read()
        for var in os.environ:
            config_content = config_content.replace(f"${{{var}}}", os.environ[var])
        config_dict = yaml.safe_load(config_content)
        
        # Set firm-specific database name and case-specific container
        database_name = f"graphrag_{firm_id}"
        config_dict["output"]["base_dir"] = database_name
        config_dict["output"]["container_name"] = f"output_{case_id}"
        
        # Make vector base container case-specific directly; avoid runtime patching
        if "vector_store" in config_dict and "default_vector_store" in config_dict["vector_store"]:
            config_dict["vector_store"]["default_vector_store"]["container_name"] = f"output_{case_id}"
        
        # Note: No runtime patching of create_index_name during query; using case-specific base container above
        
        # CRITICAL: Set database_name and connection_string for vector store (required for CosmosDB vector store)
        config_dict["vector_store"]["default_vector_store"]["database_name"] = database_name
        
        # Explicitly set connection_string to ensure AAD fallback doesn't happen
        connection_string = os.getenv("COSMOS_CONNECTION_STRING")
        if not connection_string:
            # Build from endpoint and key if connection_string not available
            endpoint = os.getenv("COSMOS_ENDPOINT")
            key = os.getenv("COSMOS_KEY")
            if endpoint and key:
                connection_string = f"AccountEndpoint={endpoint};AccountKey={key};"
        
        if connection_string:
            config_dict["vector_store"]["default_vector_store"]["connection_string"] = connection_string
        else:
            raise ValueError("COSMOS_CONNECTION_STRING or COSMOS_ENDPOINT+COSMOS_KEY must be set")
        
        try:
            # Note: db_uri is only needed for lancedb, not for cosmosdb vector store
            # Since we're using cosmosdb, we don't set db_uri
            
            config_obj = create_graphrag_config(values=config_dict, root_dir=RAG_ROOT)
            
            # Load graph data from CosmosDB using GraphRAG's native storage utility
            print(f"[COSMOS] Loading graph data from CosmosDB for case {case_id}...")
            
            entities = await firm_output_manager.get_graph_data_as_dataframe(case_id, "entities", firm_id)
            communities = await firm_output_manager.get_graph_data_as_dataframe(case_id, "communities", firm_id)
            community_reports = await firm_output_manager.get_graph_data_as_dataframe(case_id, "community_reports", firm_id)
            text_units = await firm_output_manager.get_graph_data_as_dataframe(case_id, "text_units", firm_id)
            relationships = await firm_output_manager.get_graph_data_as_dataframe(case_id, "relationships", firm_id)
            
            if entities is None:
                raise ValueError(f"No indexed data found for case {case_id}")
            
            # Covariates are optional - try loading as DataFrame
            try:
                covariates = await firm_output_manager.get_graph_data_as_dataframe(case_id, "covariates", firm_id)
            except Exception:
                # Fallback to single document approach if not found
                covariates_doc = firm_output_manager.get_graph_data(case_id, "covariates")
                covariates = pd.DataFrame(covariates_doc["data"]) if covariates_doc and "data" in covariates_doc else None
            
            # Perform local search
            response, context = await local_search(
                config_obj,
                entities,
                communities,
                community_reports,
                text_units,
                relationships,
                covariates,
                community_level=1,
                response_type="default",
                query=query_text
            )
            
            print(f"Query executed for case_id={case_id}")
            print(f"Answer: {response}")
            firm_logs_manager.info(case_id, f"Query successful: {query_text[:50]}...", "query")
            
            return {"response": response, "context": context}
        
        finally:
            pass
    
    except Exception as e:
        print(f"GraphRAG query failed: {e}")
        logs_manager.error(case_id, f"Query failed: {e}", "query")
        return {"response": "", "context": []}
